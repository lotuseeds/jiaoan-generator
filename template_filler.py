"""
Word 模板填充模块
将AI生成的内容填入模板，保留原有格式
教学方案（t4）支持：左栏知识内容 + 图片（紧跟子节）；右栏教学活动（含思政）
"""
import re
import os
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH

# 左栏标题层级正则
_L1_RE = re.compile(          # 一级标题：顶格加粗
    r'^(【.+?】|[一二三四五六七八九十]+[、.．])'
)
_L2_RE = re.compile(          # 二级标题：顶格加粗
    r'^([※△]\s*（[一二三四五六七八九十]+）|（[一二三四五六七八九十]+）)'
)
_L3_RE = re.compile(          # 三级标题：1. 2. 3.（加粗+灰色底纹）
    r'^\d+[\.．]\s*\S'
)
# 四级及以下、正文：首行缩进2个全角字符，不加粗

# 【活动标签】和英文单词的正则（用于灵动化格式）
_LABEL_RE = re.compile(r'(【[^】]+】)')
_ENGLISH_RE = re.compile(r'([A-Za-z][A-Za-z0-9\-\.]*)')
_TIME_RE = re.compile(r'^时间[：:]\s*\d+\s*(?:min|分钟)', re.IGNORECASE)
_SEP_RE  = re.compile(r'^-{3,}\s*$')

# 阶段标签：不触发"模块描述行"斜体逻辑
_PHASE_LABELS = {'导课', '正文', '小结', '课后布置', '讲授', '课堂小结'}


def _add_gray_shading(para):
    """为段落添加灰色底纹（用于三级标题）"""
    pPr = para._element.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'D9D9D9')
    pPr.append(shd)


def _strip_markdown(line: str) -> str:
    """去除 AI 可能输出的 markdown 残留（行首 #、行内 **、*、_ 标记）"""
    line = re.sub(r'^#+\s*', '', line)
    line = re.sub(r'\*\*([^*]+)\*\*', r'\1', line)
    line = re.sub(r'\*([^*]+)\*', r'\1', line)
    line = re.sub(r'_([^_]+)_', r'\1', line)
    return line


def _write_vivid_to_para(para, line: str, base_bold: bool = False):
    """
    将一行文字写入段落，自动应用灵动格式：
    - 【活动标签】→ 加粗 + 下划线
    - 英文单词/术语 → 斜体
    - 其余内容 → 正常（或随 base_bold 加粗）
    """
    parts = _LABEL_RE.split(line)
    for part in parts:
        if not part:
            continue
        if _LABEL_RE.fullmatch(part):
            run = para.add_run(part)
            run.font.bold = True
            run.font.underline = True
        else:
            subparts = _ENGLISH_RE.split(part)
            for sub in subparts:
                if not sub:
                    continue
                run = para.add_run(sub)
                run.font.bold = base_bold
                if _ENGLISH_RE.fullmatch(sub):
                    run.font.italic = True


def _is_title_line(line: str) -> bool:
    """判断一行文字是否为标题行（需要加粗）"""
    return bool(_TITLE_PATTERNS.match(line.strip()))


def _clear_cell(cell):
    """清空单元格内容但保留格式"""
    for para in cell.paragraphs:
        for run in para.runs:
            run.text = ""
    for para in cell.paragraphs[1:]:
        p = para._element
        p.getparent().remove(p)


_MD_RE = re.compile(r'(\*\*[^*\n]+\*\*|\*[^*\n]+\*)')

def _write_markdown_to_cell(cell, text: str):
    """将含 markdown 格式的文本写入单元格：**加粗** → 粗体，*斜体* → 斜体，其余原样输出"""
    _clear_cell(cell)
    for line in text.split("\n"):
        para = cell.add_paragraph()
        parts = _MD_RE.split(line)
        for part in parts:
            if not part:
                continue
            if part.startswith("**") and part.endswith("**") and len(part) > 4:
                run = para.add_run(part[2:-2])
                run.bold = True
            elif part.startswith("*") and part.endswith("*") and len(part) > 2:
                run = para.add_run(part[1:-1])
                run.italic = True
            else:
                para.add_run(part)


def _write_director_comment(cell, comment: str):
    """写入主任批语：仿宋、三号(16pt)、1.5倍行距、首行缩进2全角字符（32pt）"""
    try:
        # 在模板第2段（标题行后第一个空段落）写入，保留模板原有结构
        para = cell.paragraphs[1] if len(cell.paragraphs) > 1 else cell.add_paragraph()
        for run in para.runs:
            run.text = ""
        para.paragraph_format.first_line_indent = Pt(32)
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        run = para.add_run(comment)
        run.font.name = "仿宋"
        _set_east_asia_font(run, "仿宋")
        run.font.size = Pt(16)

        # 删除多余空行：批语后只保留1个空行，然后是签名行
        # 目标结构：[0]标题 [1]批语 [2]空行 [3]签名
        while len(cell.paragraphs) > 4:
            candidate = cell.paragraphs[3]
            if candidate.text.strip():  # 遇到有内容的段落（签名行），停止
                break
            candidate._element.getparent().remove(candidate._element)
    except Exception:
        pass


def _set_cell_valign(cell, val: str = "center"):
    """设置单元格垂直对齐（top/center/bottom）"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement("w:vAlign")
    vAlign.set(qn("w:val"), val)
    tcPr.append(vAlign)


def _set_cell_rich(cell, text: str, font_name: str = None, font_size: float = None,
                   bold: bool = False, alignment=None):
    """设置单元格文字，支持字体名（含中文字体）、字号、加粗、对齐"""
    _clear_cell(cell)
    para = cell.paragraphs[0]
    if alignment is not None:
        para.paragraph_format.alignment = alignment
    lines = text.split("\n")
    for i, line in enumerate(lines):
        if i == 0:
            run = para.add_run(line)
        else:
            para = cell.add_paragraph()
            if alignment is not None:
                para.paragraph_format.alignment = alignment
            run = para.add_run(line)
        if font_name:
            run.font.name = font_name
            _set_east_asia_font(run, font_name)
        if font_size is not None:
            run.font.size = Pt(font_size)
        run.font.bold = bold


def _set_cell_text(cell, text: str, font_size=None, bold=False):
    """设置单元格文字（清空后写入），保留原有字体格式"""
    _clear_cell(cell)
    para = cell.paragraphs[0]
    lines = text.split("\n")
    for i, line in enumerate(lines):
        if i == 0:
            run = para.add_run(line)
        else:
            para = cell.add_paragraph()
            run = para.add_run(line)
        if font_size:
            run.font.size = Pt(font_size)
        if bold:
            run.font.bold = bold


def _append_text_to_cell(cell, text: str, font_size=None, bold=False):
    """向单元格追加文字（不清空已有内容），按行分段落，自动应用灵动格式"""
    lines = text.split("\n")
    for line in lines:
        line = _strip_markdown(line)
        para = cell.add_paragraph()
        is_title = _is_title_line(line)
        _write_vivid_to_para(para, line, base_bold=(bold or is_title))
        if font_size:
            for run in para.runs:
                run.font.size = Pt(font_size)


def _append_right_col_to_cell(cell, text: str, font_size=None):
    """右栏专用写入：
    - 时间：X min      → 小四(12pt) + 加粗 + 斜体 + 下划线
    - 时间行下一非空行  → 关键词行，斜体，不加下划线
    - 【模块名】        → 加粗，不加下划线
    - 【模块名】后第一段 → 斜体
    - ---              → 渲染为横线
    - 其余内容          → 正常
    """
    after_time = False    # 下一个非空行是关键词行
    after_module = False  # 下一个非空行是模块描述行

    for line in text.split("\n"):
        line_clean = _strip_markdown(line)
        stripped = line_clean.strip()
        para = cell.add_paragraph()

        if _SEP_RE.match(line.strip()):
            para.add_run("─" * 24)
            after_time = False
            after_module = False

        elif _TIME_RE.match(line.strip()):
            run = para.add_run(line_clean)
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.italic = True
            run.font.underline = True
            after_time = True
            after_module = False

        elif stripped.startswith('【'):
            # 模块标题行：加粗，不加下划线
            parts = _LABEL_RE.split(line_clean)
            for part in parts:
                if not part:
                    continue
                run = para.add_run(part)
                run.font.bold = bool(_LABEL_RE.fullmatch(part))
            if font_size:
                for run in para.runs:
                    run.font.size = Pt(font_size)
            # 阶段标签（【导课】【小结】等）不触发描述行逻辑
            m = re.match(r'【([^】]+)】', stripped)
            is_phase = bool(m and m.group(1) in _PHASE_LABELS)
            after_module = not is_phase
            after_time = False

        elif after_time and stripped:
            # 时间行下方的关键词行：小五(9pt) + 斜体
            run = para.add_run(line_clean)
            run.font.italic = True
            run.font.size = Pt(9)
            after_time = False

        elif after_module and stripped:
            # 模块名后的第一段描述：小五(9pt) + 斜体
            run = para.add_run(line_clean)
            run.font.italic = True
            run.font.size = Pt(9)
            after_module = False

        else:
            if stripped:
                after_time = False
                after_module = False
            _write_vivid_to_para(para, line_clean, base_bold=False)
            for run in para.runs:
                run.font.size = Pt(font_size) if font_size else Pt(10.5)  # 五号


def _append_left_col_to_cell(cell, text: str, font_size=None, shading=True):
    """左栏专用写入（1.5倍行距）：
    - 一级标题（【】或一、二、三...）→ 顶格加粗，四号(14pt)
    - 二级标题（※（一）△（二）等）  → 顶格加粗，四号(14pt)
    - 三级标题（1. 2. 3.）          → 缩进加粗，小四(12pt)，灰色底纹（shading=False时不加）
    - 四级及以下 + 正文              → 缩进不加粗，小四(12pt)
    """
    for line in text.split("\n"):
        line_clean = _strip_markdown(line)
        stripped = line_clean.strip()
        para = cell.add_paragraph()
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

        if _L1_RE.match(stripped) or _L2_RE.match(stripped):
            # 一级、二级：顶格加粗，四号(14pt)
            _write_vivid_to_para(para, stripped, base_bold=True)
            for run in para.runs:
                run.font.size = Pt(14)
        elif _L3_RE.match(stripped):
            # 三级标题判断模式：
            # - 正文（shading=True）：≤25字 → 标题形式（加粗+灰色底纹），否则段落形式
            # - 导课/小结/课后（shading=False）：始终段落形式（不加粗、不加底纹）
            is_heading_form = shading and len(stripped) <= 25
            if is_heading_form:
                _write_vivid_to_para(para, "　　" + stripped, base_bold=True)
                for run in para.runs:
                    run.font.size = Pt(12)
                _add_gray_shading(para)
            else:
                _write_vivid_to_para(para, "　　" + stripped, base_bold=False)
                for run in para.runs:
                    run.font.size = Pt(12)
        else:
            # 四级及以下、正文：缩进不加粗，小四(12pt)
            display = ("　　" + stripped) if stripped else ""
            _write_vivid_to_para(para, display, base_bold=False)
            for run in para.runs:
                run.font.size = Pt(12)


def _split_at_l3_boundaries(text: str) -> list:
    """
    将内容文本按标题层级切分为若干 segment：
    - 遇到 L3 标题（1. 2. 3.）：新 segment，l3_title=该标题
    - 遇到 L1/L2 标题：新 segment，l3_title=None（防止 L2 标题漏进前一个 L3 段）
    返回：[{"l3_title": None/"1. xxx", "text": "..."}, ...]
    """
    segments = []
    current_title = None
    current_lines = []

    for line in text.split("\n"):
        stripped = _strip_markdown(line).strip()
        is_l1 = bool(_L1_RE.match(stripped))
        is_l2 = bool(_L2_RE.match(stripped))
        is_l3 = bool(_L3_RE.match(stripped)) and not is_l1 and not is_l2

        if is_l3:
            if current_lines:
                segments.append({"l3_title": current_title,
                                  "text": "\n".join(current_lines)})
            current_title = stripped
            current_lines = [line]
        elif is_l1 or is_l2:
            # L1/L2 另起 segment，重置 l3_title，防止 L2 漏进上一个 L3 段
            if current_lines:
                segments.append({"l3_title": current_title,
                                  "text": "\n".join(current_lines)})
            current_title = None
            current_lines = [line]
        else:
            current_lines.append(line)

    if current_lines:
        segments.append({"l3_title": current_title,
                          "text": "\n".join(current_lines)})
    return segments


def _set_east_asia_font(run, font_name: str):
    """设置 run 的中文字体（东亚字体）"""
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)


def _add_image_to_cell(cell, image_path: str, width_inches=2.5):
    """在单元格末尾插入图片（居中）"""
    try:
        para = cell.add_paragraph()
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        run.add_picture(image_path, width=Inches(width_inches))
    except Exception:
        para = cell.add_paragraph()
        para.add_run(f"[图片: {os.path.basename(image_path)}]")


def _write_chapter_title(cell, title: str):
    """写入正文章节标题：三号(16pt) 黑体 居中 加粗 1.5倍行距"""
    para = cell.add_paragraph()
    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = para.add_run(title)
    run.font.bold = True
    run.font.size = Pt(16)  # 三号
    run.font.name = '黑体'
    _set_east_asia_font(run, '黑体')


def _fill_teaching_plan_cell(
    left_cell,
    right_cell,
    plan_items: list,
    slide_image_map: dict,
    max_images_per_section: int = 2,
    chapter_title: str = "",
):
    """
    逐段落写入教学方案，在每个子节文字之后插入对应图片。

    plan_items 结构（来自 ai_content["teaching_plan"]）：
      每个 item 可以有 sections（正文）或没有（导课/小结/课后）
      sections 中每个子节有 content/activity/related_slides

    左栏（left_cell）：知识内容 + 图片（紧跟子节之后）
    右栏（right_cell）：教学活动（含思政、互动、手段）
    chapter_title：授课章节/主题，用于替换【正文】标签
    """
    _clear_cell(left_cell)
    _clear_cell(right_cell)

    for item in plan_items:
        phase = item.get("phase", "")
        duration = item.get("duration", 0)
        sections = item.get("sections")  # 正文才有

        if sections:
            # ── 正文：用章节标题替换【正文】──
            if chapter_title:
                _write_chapter_title(left_cell, chapter_title)
            else:
                _append_left_col_to_cell(left_cell, f"【{phase}】")
            _append_right_col_to_cell(right_cell, f"【{phase}】（共{duration}min）")

            for sec in sections:
                sec_content = sec.get("content", "")
                sec_activity = sec.get("activity", "")
                image_assignments = sec.get("image_assignments", [])

                if image_assignments:
                    # 按三级标题分段写入，在匹配的三级标题内容后插图
                    assign_map = {a["l3_title"]: a["slide_index"]
                                  for a in image_assignments}
                    for seg in _split_at_l3_boundaries(sec_content):
                        _append_left_col_to_cell(left_cell, seg["text"])
                        l3 = seg.get("l3_title")
                        if l3 and l3 in assign_map:
                            img_path = slide_image_map.get(assign_map[l3])
                            if img_path:
                                _add_image_to_cell(left_cell, img_path,
                                                   width_inches=2.5)
                else:
                    # 降级：整节写完再插图（Stage 3 未运行或无匹配截图）
                    _append_left_col_to_cell(left_cell, sec_content)
                    inserted = 0
                    for slide_idx in sec.get("related_slides", []):
                        if inserted >= max_images_per_section:
                            break
                        img_path = slide_image_map.get(slide_idx)
                        if img_path:
                            _add_image_to_cell(left_cell, img_path,
                                               width_inches=2.5)
                            inserted += 1

                # 右栏：写子节教学活动
                _append_right_col_to_cell(right_cell, sec_activity)

        else:
            # ── 导课 / 小结 / 课后：单块写入 ──
            content = item.get("content", "")
            activity = item.get("activity", "")
            related = item.get("related_slides", [])

            _append_left_col_to_cell(left_cell, f"【{phase}】\n{content}", shading=False)

            # 左栏：插入该阶段对应图片
            inserted = 0
            for slide_idx in related:
                if inserted >= max_images_per_section:
                    break
                img_path = slide_image_map.get(slide_idx)
                if img_path:
                    _add_image_to_cell(left_cell, img_path, width_inches=2.5)
                    inserted += 1

            _append_right_col_to_cell(right_cell, f"【{phase}】\n{activity}")


def fill_template(
    template_path: str,
    output_path: str,
    basic_info: dict,
    ai_content: dict,
    ppt_data: dict = None,
    max_images_per_section: int = 2,
):
    """
    填充模板并输出新文件

    basic_info: 基本信息字典
    ai_content: AI生成的完整教案内容
    ppt_data:   完整PPT解析数据（用于按页码查找图片）
    max_images_per_section: 每个知识子节最多插入几张图片
    """
    doc = Document(template_path)
    tables = doc.tables

    # 构建 slide_index → image_path 映射
    slide_image_map = {}
    if ppt_data:
        for slide in ppt_data.get("slides", []):
            if slide.get("has_image") and slide.get("image_path"):
                slide_image_map[slide["index"]] = slide["image_path"]

    # ── 表格1：封面基本信息 ──
    t1 = tables[0]
    _t1_fields = [
        basic_info.get("course_name", ""),
        basic_info.get("teacher_name", ""),
        basic_info.get("professional_title", ""),
        basic_info.get("department", ""),
        basic_info.get("college", ""),
        basic_info.get("students", ""),
    ]
    for row_idx, val in enumerate(_t1_fields):
        cell = t1.cell(row_idx, 1)
        _set_cell_rich(cell, val, font_name="楷体", font_size=15,
                       alignment=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_valign(cell, "center")

    # ── 表格3：教学计划 ──
    t3 = tables[2]

    # 教学内容：楷体、四号(14pt)、加粗
    _set_cell_rich(t3.cell(0, 2), basic_info.get("title", ""),
                   font_name="楷体", font_size=14, bold=True,
                   alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # 教材名称：加《》，仿宋、小四(12pt)、居中
    book_name = basic_info.get("textbook_name", "")
    if book_name and not book_name.startswith("《"):
        book_name = f"《{book_name}》"
    textbook_info = (
        f"{book_name}\n"
        f"{basic_info.get('textbook_publisher', '')}\n"
        f"{basic_info.get('textbook_edition', '')}"
    )
    for row_idx in [1, 2, 3]:
        try:
            _set_cell_rich(t3.cell(row_idx, 2), textbook_info,
                           font_name="仿宋", font_size=12,
                           alignment=WD_ALIGN_PARAGRAPH.CENTER)
        except Exception:
            pass

    # 主编、教材性质、出版时间：仿宋、小四(12pt)、居中
    try:
        _set_cell_rich(t3.cell(1, 4), basic_info.get("textbook_editor", ""),
                       font_name="仿宋", font_size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_rich(t3.cell(2, 4), basic_info.get("textbook_series", ""),
                       font_name="仿宋", font_size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_rich(t3.cell(3, 4), basic_info.get("textbook_year", ""),
                       font_name="仿宋", font_size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    except Exception:
        pass

    # 教学时间：仿宋、小四(12pt)、居中
    try:
        _set_cell_rich(t3.cell(4, 2), basic_info.get("teaching_date", "") + "，计  2学时",
                       font_name="仿宋", font_size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    except Exception:
        pass

    # 教学对象、教学地点：仿宋、小四(12pt)、居中
    try:
        _set_cell_rich(t3.cell(6, 2), basic_info.get("students", ""),
                       font_name="仿宋", font_size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_rich(t3.cell(6, 4), basic_info.get("classroom", ""),
                       font_name="仿宋", font_size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    except Exception:
        pass

    try:
        _set_cell_text(t3.cell(8, 2), ai_content.get("teaching_objective_knowledge", ""))
        _set_cell_text(t3.cell(9, 2), ai_content.get("teaching_objective_ability", ""))
        _set_cell_text(t3.cell(10, 2), ai_content.get("teaching_objective_value", ""))
    except Exception:
        pass

    try:
        _set_cell_text(t3.cell(11, 2), ai_content.get("student_analysis_knowledge", ""))
        _set_cell_text(t3.cell(12, 2), ai_content.get("student_analysis_cognition", ""))
        _set_cell_text(t3.cell(13, 2), ai_content.get("student_analysis_psychology", ""))
    except Exception:
        pass

    try:
        _set_cell_text(t3.cell(14, 2), ai_content.get("key_points", ""))
        _set_cell_text(t3.cell(15, 2), ai_content.get("difficult_points", ""))
    except Exception:
        pass

    try:
        _set_cell_text(t3.cell(16, 2), ai_content.get("teaching_expansion", ""))
    except Exception:
        pass

    try:
        _set_cell_text(t3.cell(17, 2), ai_content.get("self_study", ""))
    except Exception:
        pass

    try:
        _set_cell_text(t3.cell(18, 2), ai_content.get("teaching_method", ""))
        _set_cell_text(t3.cell(19, 2), ai_content.get("teaching_tools", ""))
        _set_cell_text(t3.cell(20, 2), ai_content.get("keywords_en", ""))
        _set_cell_text(t3.cell(21, 2), ai_content.get("references", ""))
    except Exception:
        pass

    # ── 表格4：教学方案（两栏：内容架构 + 教学活动）──
    t4 = tables[3]
    plan_items = ai_content.get("teaching_plan", [])

    try:
        _fill_teaching_plan_cell(
            left_cell=t4.cell(1, 0),
            right_cell=t4.cell(1, 1),
            plan_items=plan_items,
            slide_image_map=slide_image_map,
            max_images_per_section=max_images_per_section,
            chapter_title=basic_info.get("title", ""),
        )
    except Exception:
        pass

    # ── 表格5：主板书设计 + 课后习题 ──
    t5 = tables[4]
    try:
        _set_cell_text(t5.cell(0, 1), ai_content.get("blackboard_left", ""))
        _set_cell_text(t5.cell(0, 2), ai_content.get("blackboard_right", ""))
        _write_markdown_to_cell(t5.cell(1, 1), ai_content.get("homework", ""))
    except Exception:
        pass

    # ── 表格6：自主学习资源 ──
    t6 = tables[5]
    try:
        _write_markdown_to_cell(t6.cell(0, 2), ai_content.get("self_study_resources", ""))
    except Exception:
        pass

    # ── 主任批语 ──
    try:
        director_comment = ai_content.get("director_comment", "")
        if director_comment:
            _write_director_comment(tables[5].cell(1, 0), director_comment)
    except Exception:
        pass

    doc.save(output_path)
    return output_path
